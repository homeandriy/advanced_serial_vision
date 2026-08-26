"""Build the Ukrainian Serial Vision user guide from the in-app help source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).parents[1]
SOURCE = ROOT / "application" / "serial_vision" / "assets" / "help" / "uk.md"
OUTPUT = ROOT / "documents" / "serial-vision-user-guide.uk.docx"


def set_font(run, name: str, size: float, color: str | None = None, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)


def parse_sections(markdown: str) -> list[tuple[str, list[str]]]:
    chunks = re.split(r"<!-- section: [a-z-]+ -->\s*", markdown)
    sections: list[tuple[str, list[str]]] = []
    for chunk in chunks:
        lines = [line.strip() for line in chunk.strip().splitlines()]
        if not lines:
            continue
        title = lines[0].removeprefix("## ")
        paragraphs = [line for line in lines[1:] if line]
        sections.append((title, paragraphs))
    return sections


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 18, "1F4D78", 18, 8),
        ("Heading 2", 16, "2E74B5", 14, 6),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("Serial Vision | Посібник користувача")
    set_font(header_run, "Aptos", 9, "5B6573")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Сторінка ")
    set_font(footer_run, "Aptos", 9, "5B6573")
    add_page_field(footer)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(86)
    title.paragraph_format.space_after = Pt(12)
    set_font(title.add_run("SERIAL VISION"), "Aptos Display", 28, "1F4D78", True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(subtitle.add_run("Розширений посібник користувача"), "Aptos", 18, "2E74B5")
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.left_indent = Cm(1.2)
    intro.paragraph_format.right_indent = Cm(1.2)
    intro.paragraph_format.line_spacing = 1.2
    set_font(intro.add_run("Локальний облік обладнання, розпізнавання серійних номерів, MAC-адрес і штрихкодів"), "Aptos", 14, "394B59")
    document.add_page_break()

    extra_source = SOURCE.with_name("uk-extra.md")
    sections = parse_sections(SOURCE.read_text(encoding="utf-8") + extra_source.read_text(encoding="utf-8"))
    document.add_heading("Зміст", level=1)
    for index, (title_text, _) in enumerate(sections, start=1):
        item = document.add_paragraph(style="Normal")
        item.paragraph_format.left_indent = Cm(0.5)
        item.paragraph_format.first_line_indent = Cm(-0.5)
        set_font(item.add_run(f"{index}. {title_text}"), "Aptos", 14, "1F4D78", True)
    document.add_page_break()

    for index, (title_text, paragraphs) in enumerate(sections, start=1):
        document.add_heading(f"{index}. {title_text}", level=1)
        for paragraph_text in paragraphs:
            paragraph = document.add_paragraph(style="Normal")
            paragraph.paragraph_format.first_line_indent = Cm(0.7)
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_font(paragraph.add_run(paragraph_text), "Aptos", 14, "202124")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = "Serial Vision — Розширений посібник користувача"
    document.core_properties.subject = "Довідка з використання Serial Vision"
    document.core_properties.author = "Serial Vision"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
